"""Tests for document generator — DOCX and PDF output."""
import pytest
from io import BytesIO
from docx import Document as DocxDocument


@pytest.fixture
def sample_html():
    return '<p>Ο Αντιπεριφερειάρχης Κοινωνικής Μέριμνας αποφασίζει:</p><p>Χορηγείται άδεια λειτουργίας.</p>'


@pytest.fixture
def sample_title():
    return 'Απόφαση Λειτουργίας Κατασκήνωσης'


@pytest.fixture
def sample_recipients():
    return [
        {'name': 'Ενδιαφερόμενος'},
        {'name': 'Δ/νση Κοινωνικής Μέριμνας'},
    ]


@pytest.fixture
def sample_legal_refs():
    return [
        "Ν.4939/2022 (ΦΕΚ Α' 111)",
        "Ν.3852/2010 (ΦΕΚ Α' 87)",
    ]


class TestGenerateDecisionDocx:
    def test_returns_bytes(self, sample_html, sample_title, sample_recipients, sample_legal_refs):
        from my_project.documents.generator import generate_decision_docx
        result = generate_decision_docx(sample_html, sample_title, sample_recipients, sample_legal_refs)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_valid_docx_format(self, sample_html, sample_title, sample_recipients, sample_legal_refs):
        from my_project.documents.generator import generate_decision_docx
        result = generate_decision_docx(sample_html, sample_title, sample_recipients, sample_legal_refs)
        # python-docx can read it back
        doc = DocxDocument(BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_contains_title(self, sample_html, sample_title, sample_recipients, sample_legal_refs):
        from my_project.documents.generator import generate_decision_docx
        result = generate_decision_docx(sample_html, sample_title, sample_recipients, sample_legal_refs)
        doc = DocxDocument(BytesIO(result))
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'ΘΕΜΑ' in all_text
        assert sample_title in all_text

    def test_contains_body_text(self, sample_html, sample_title, sample_recipients, sample_legal_refs):
        from my_project.documents.generator import generate_decision_docx
        result = generate_decision_docx(sample_html, sample_title, sample_recipients, sample_legal_refs)
        doc = DocxDocument(BytesIO(result))
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'Αντιπεριφερειάρχης' in all_text

    def test_contains_legal_references(self, sample_html, sample_title, sample_recipients, sample_legal_refs):
        from my_project.documents.generator import generate_decision_docx
        result = generate_decision_docx(sample_html, sample_title, sample_recipients, sample_legal_refs)
        doc = DocxDocument(BytesIO(result))
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'Έχοντας υπόψη' in all_text
        assert 'Ν.4939/2022' in all_text

    def test_contains_recipients(self, sample_html, sample_title, sample_recipients, sample_legal_refs):
        from my_project.documents.generator import generate_decision_docx
        result = generate_decision_docx(sample_html, sample_title, sample_recipients, sample_legal_refs)
        doc = DocxDocument(BytesIO(result))
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        table_text = ''
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + ' '
        combined = all_text + table_text
        assert 'ΑΠΟΔΕΚΤ' in combined.upper()

    def test_a4_page_dimensions(self, sample_html, sample_title, sample_recipients, sample_legal_refs):
        from my_project.documents.generator import generate_decision_docx
        result = generate_decision_docx(sample_html, sample_title, sample_recipients, sample_legal_refs)
        doc = DocxDocument(BytesIO(result))
        section = doc.sections[0]
        # A4 = 210mm x 297mm in EMU
        assert abs(section.page_width - 7560310) < 50000
        assert abs(section.page_height - 10692130) < 50000

    def test_contains_header_block(self, sample_html, sample_title, sample_recipients, sample_legal_refs):
        from my_project.documents.generator import generate_decision_docx
        result = generate_decision_docx(sample_html, sample_title, sample_recipients, sample_legal_refs)
        doc = DocxDocument(BytesIO(result))
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        table_text = ''
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + '\n'
        combined = all_text + table_text
        assert 'ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ' in combined

    def test_without_legal_refs(self, sample_html, sample_title, sample_recipients):
        from my_project.documents.generator import generate_decision_docx
        result = generate_decision_docx(sample_html, sample_title, sample_recipients, legal_references=None)
        assert isinstance(result, bytes)
        doc = DocxDocument(BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_without_recipients(self, sample_html, sample_title, sample_legal_refs):
        from my_project.documents.generator import generate_decision_docx
        result = generate_decision_docx(sample_html, sample_title, recipients=None, legal_references=sample_legal_refs)
        assert isinstance(result, bytes)


class TestGenerateDecisionPdf:
    """Verify existing PDF generation still works."""
    def test_returns_bytes(self, sample_html, sample_title, sample_recipients):
        from my_project.documents.generator import generate_decision_pdf
        result = generate_decision_pdf(sample_html, sample_title, sample_recipients)
        assert isinstance(result, bytes)
        assert result[:4] == b'%PDF'
