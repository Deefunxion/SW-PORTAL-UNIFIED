"""
One-time script to extract evaluation criteria from Ministerial Decision
.docx templates in content/ΕΚΘΕΣΕΙΣ_ΕΛΕΓΧΩΝ/
Outputs structured JSON that maps to INSPECTION_CRITERIA format.
"""
import json
import os
import sys

try:
    from docx import Document
except ImportError:
    print("Install python-docx: pip install python-docx")
    sys.exit(1)

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), '..', 'content', 'ΕΚΘΕΣΕΙΣ_ΕΛΕΓΧΩΝ')


def extract_docx(path):
    """Extract text content and table structure from a .docx file."""
    doc = Document(path)
    result = {'paragraphs': [], 'tables': []}
    for para in doc.paragraphs:
        if para.text.strip():
            result['paragraphs'].append({
                'text': para.text.strip(),
                'style': para.style.name if para.style else None,
                'bold': any(run.bold for run in para.runs),
            })
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_data.append([cell.text.strip() for cell in row.cells])
        result['tables'].append(table_data)
    return result


def main():
    output = {}
    for root, dirs, files in os.walk(TEMPLATE_DIR):
        for f in files:
            if f.endswith('.docx'):
                path = os.path.join(root, f)
                print(f"Extracting: {f}")
                try:
                    output[f] = extract_docx(path)
                except Exception as e:
                    print(f"  Error: {e}")

    out_path = os.path.join(os.path.dirname(__file__), 'template_structure.json')
    with open(out_path, 'w', encoding='utf-8') as fp:
        json.dump(output, fp, ensure_ascii=False, indent=2)
    print(f"\nOutput: {out_path}")


if __name__ == '__main__':
    main()
