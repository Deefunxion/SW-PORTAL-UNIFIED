import fitz  # PyMuPDF
import docx
import re
import spacy
import os

# --- Model Loading ---
# Load spaCy models. It's recommended to have these downloaded beforehand.
# python -m spacy download el_core_news_sm
# python -m spacy download en_core_web_sm
try:
    nlp_el = spacy.load("el_core_news_sm")
except OSError:
    print("Warning: Greek spaCy model 'el_core_news_sm' not found. Greek NER will be skipped.")
    nlp_el = None

try:
    nlp_en = spacy.load("en_core_web_sm")
except OSError:
    print("Warning: English spaCy model 'en_core_web_sm' not found. English NER will be skipped.")
    nlp_en = None

# --- PII Definitions ---
# Comprehensive regex for various phone number formats
phone_regex = r'(\+?\d{1,3}[\s.-]?)?\(?\d{2,4}\)?[\s.-]?\d{3,4}[\s.-]?\d{3,4}'

# Regex patterns for structured Greek PII
PII_PATTERNS = {
    'GREEK_AFM': r'\b\d{9}\b',
    'GREEK_ADT': r'\b[Α-Ωα-ω]{2}\s?\d{6}\b',
    'GREEK_AMKA': r'\b\d{11}\b',
    'PHONE_NUMBER': phone_regex,
    'EMAIL': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
}

# --- Text Extraction ---
def extract_text_from_pdf(file_path):
    """Extracts all text from a PDF file."""
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_path):
    """Extracts all text from a DOCX file."""
    text = ""
    doc = docx.Document(file_path)
    for para in doc.paragraphs:
        text += para.text + "\n"
    # You might want to extend this to handle tables, headers, etc.
    return text

# --- PII Detection ---
def find_pii_in_text(text):
    """Finds all PII instances in a given text using regex and NER."""
    found_pii = []

    # 1. Regex-based search for structured PII
    for pii_type, pattern in PII_PATTERNS.items():
        matches = re.finditer(pattern, text, flags=re.IGNORECASE)
        for match in matches:
            found_pii.append(match.group(0))

    # 2. NER-based search for unstructured PII (Names, Locations)
    if nlp_el:
        doc_el = nlp_el(text)
        for ent in doc_el.ents:
            if ent.label_ in ['PERSON', 'GPE', 'LOC']:
                found_pii.append(ent.text)
    
    if nlp_en:
        doc_en = nlp_en(text)
        for ent in doc_en.ents:
            if ent.label_ in ['PERSON', 'GPE', 'LOC', 'ORG']: # ORG can sometimes catch professions/institutions
                found_pii.append(ent.text)

    # Return a list of unique PII strings, prioritizing longer strings
    # to avoid redacting "John" if "John Smith" is also found.
    unique_pii = sorted(list(set(found_pii)), key=len, reverse=True)
    return unique_pii

# --- Redaction Logic ---
def redact_pdf(file_path, pii_to_redact):
    """Redacts specified PII from a PDF file."""
    doc = fitz.open(file_path)
    for page in doc:
        for pii in pii_to_redact:
            # Search for the text and get redaction areas
            areas = page.search_for(pii, quads=True)
            for area in areas:
                page.add_redact_annot(area, fill=(0, 0, 0)) # Black fill
    
    # Apply all redactions at once
    doc.apply_redactions()
    # Save the redacted file, overwriting the original
    doc.save(file_path)
    doc.close()

def redact_docx(file_path, pii_to_redact):
    """Redacts specified PII from a DOCX file."""
    doc = docx.Document(file_path)
    for pii in pii_to_redact:
        # Build a regex to find the whole word to avoid partial matches
        # e.g., don't redact 'cat' in 'caterpillar'
        pii_regex = r'\b' + re.escape(pii) + r'\b'
        for para in doc.paragraphs:
            if re.search(pii_regex, para.text):
                # Using runs to preserve formatting
                inline = para.runs
                # Replace strings and retain formatting
                for i in range(len(inline)):
                    if re.search(pii_regex, inline[i].text):
                        inline[i].text = re.sub(pii_regex, '[REDACTED]', inline[i].text)

    # Save the redacted file, overwriting the original
    doc.save(file_path)

# --- Main Function ---
def redact_pii_in_file(file_path):
    """
    Orchestrates the PII redaction process for a given file.
    
    Args:
        file_path (str): The absolute path to the file to be redacted.
    """
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return

    file_ext = os.path.splitext(file_path)[1].lower()
    text = ""

    print(f"Starting PII redaction for: {file_path}")

    if file_ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif file_ext == '.docx':
        text = extract_text_from_docx(file_path)
    else:
        print(f"Unsupported file type: {file_ext}. Skipping redaction.")
        return

    if not text:
        print("No text extracted. Skipping redaction.")
        return

    # Find all PII in the extracted text
    pii_to_redact = find_pii_in_text(text)

    if not pii_to_redact:
        print("No PII found. File remains unchanged.")
        return
        
    print(f"Found {len(pii_to_redact)} PII instances to redact.")

    # Apply the redaction based on file type
    if file_ext == '.pdf':
        redact_pdf(file_path, pii_to_redact)
    elif file_ext == '.docx':
        redact_docx(file_path, pii_to_redact)
        
    print(f"Redaction complete for: {file_path}")

if __name__ == '__main__':
    # Example usage for testing
    # Create a dummy file or replace with a real path
    test_file = 'test_document.docx'
    if not os.path.exists(test_file):
        doc = docx.Document()
        doc.add_paragraph("My name is ΓΕΩΡΓΙΟΣ ΠΑΠΑΔΟΠΟΥΛΟΣ and my phone is +30 2101234567.")
        doc.add_paragraph("My colleague, John Smith, has an AFM of 123456789.")
        doc.save(test_file)
        print(f"Created test file: {test_file}")

    redact_pii_in_file(os.path.abspath(test_file))