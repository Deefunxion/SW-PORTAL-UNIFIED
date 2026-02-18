"""Document template rendering engine with placeholder resolution and PDF generation.

Renders DecisionTemplate body_template by filling {{placeholder}} syntax with
data from Structure models and user-provided field values.
"""
import re
import os
import html
from io import BytesIO
from datetime import date, datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


def _register_fonts():
    """Register Greek-capable fonts. Returns font name to use."""
    font_dirs = [
        'C:/Windows/Fonts',
        '/usr/share/fonts/truetype/msttcorefonts',
        '/usr/share/fonts/truetype/dejavu',
    ]
    for d in font_dirs:
        path = os.path.join(d, 'arial.ttf')
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont('Greek', path))
                bold_path = os.path.join(d, 'arialbd.ttf')
                if os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont('Greek-Bold', bold_path))
                return 'Greek'
            except Exception:
                pass
    # Fallback to DejaVu on Linux
    deja_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
    if os.path.exists(deja_path):
        try:
            pdfmetrics.registerFont(TTFont('Greek', deja_path))
            return 'Greek'
        except Exception:
            pass
    return 'Helvetica'


def resolve_placeholders(template_body, structure, user_data):
    """
    Replace {{placeholder}} in template body with actual values.

    Args:
        template_body: HTML string with {{key}} placeholders
        structure: Structure model instance (or None)
        user_data: dict of user-provided field values

    Returns:
        Rendered HTML string
    """
    ctx = {}
    if structure:
        ctx.update({
            'όνομα_δομής': structure.name,
            'κωδικός_δομής': structure.code or '',
            'πόλη': structure.city or '',
            'οδός': structure.street or '',
            'ΤΚ': structure.postal_code or '',
            'εκπρόσωπος': structure.representative_name or '',
            'ΑΦΜ_εκπροσώπου': structure.representative_afm or '',
            'τηλέφωνο_εκπροσώπου': structure.representative_phone or '',
            'email_εκπροσώπου': structure.representative_email or '',
            'δυναμικότητα': str(structure.capacity or ''),
            'κατάσταση': structure.status or '',
            'ιδιοκτησία': structure.ownership or '',
            'αριθμός_αδείας': structure.license_number or '',
        })
        if structure.license_date:
            ctx['ημερομηνία_αδείας'] = structure.license_date.strftime('%d/%m/%Y')
        if structure.license_expiry:
            ctx['λήξη_αδείας'] = structure.license_expiry.strftime('%d/%m/%Y')
        if hasattr(structure, 'structure_type') and structure.structure_type:
            ctx['τύπος_δομής'] = structure.structure_type.name

    # Overlay user-provided data (takes precedence)
    ctx.update(user_data)

    # Format dates nicely
    for key, val in list(ctx.items()):
        if isinstance(val, (date, datetime)):
            ctx[key] = val.strftime('%d/%m/%Y')

    def replacer(match):
        key = match.group(1).strip()
        return str(ctx.get(key, f'[{key}]'))

    rendered = re.sub(r'\{\{(.+?)\}\}', replacer, template_body)
    return rendered


def generate_decision_pdf(rendered_html, title, recipients=None):
    """
    Generate a PDF from rendered HTML decision text.

    Args:
        rendered_html: HTML string (rendered template)
        title: Document title for header
        recipients: list of recipient dicts

    Returns:
        PDF bytes
    """
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=2*cm, bottomMargin=2*cm,
                            leftMargin=2.5*cm, rightMargin=2.5*cm)

    font_name = _register_fonts()

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        'GreekTitle', parent=styles['Title'],
        fontName=font_name, fontSize=14, alignment=TA_CENTER,
        spaceAfter=12))
    styles.add(ParagraphStyle(
        'GreekBody', parent=styles['Normal'],
        fontName=font_name, fontSize=11, alignment=TA_JUSTIFY,
        leading=16, spaceAfter=8))
    styles.add(ParagraphStyle(
        'GreekSmall', parent=styles['Normal'],
        fontName=font_name, fontSize=9, alignment=TA_JUSTIFY,
        leading=12, spaceAfter=4))

    story = []

    # Title
    story.append(Paragraph(title, styles['GreekTitle']))
    story.append(Spacer(1, 0.5*cm))

    # Body — split HTML into paragraphs
    clean = re.sub(r'<br\s*/?>', '\n', rendered_html)
    clean = re.sub(r'<p>', '', clean)
    clean = re.sub(r'</p>', '\n\n', clean)
    clean = re.sub(r'<[^>]+>', '', clean)
    clean = html.unescape(clean)

    for para_text in clean.split('\n\n'):
        text = para_text.strip()
        if text:
            story.append(Paragraph(text, styles['GreekBody']))

    # Recipients table
    if recipients:
        story.append(Spacer(1, 1*cm))
        story.append(Paragraph('ΠΙΝΑΚΑΣ ΑΠΟΔΕΚΤΩΝ', styles['GreekTitle']))
        for i, r in enumerate(recipients, 1):
            name = r.get('name', '')
            story.append(
                Paragraph(f'{i}. {name}', styles['GreekSmall']))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def generate_decision_docx(rendered_html, title, recipients=None,
                           legal_references=None, protocol_number=None,
                           internal_number=None, doc_date=None):
    """
    Generate a DOCX from rendered HTML decision text.

    Args:
        rendered_html: HTML string (rendered template)
        title: Document title for ΘΕΜΑ line
        recipients: list of recipient dicts [{"name": "..."}]
        legal_references: list of legal reference strings
        protocol_number: protocol number string (or None)
        internal_number: internal number string (or None)
        doc_date: date string (or None, defaults to today)

    Returns:
        DOCX bytes
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # A4 page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    font_name = 'Arial'

    # ── 1. Header table (two columns) ─────────────────────────
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left column: ministry hierarchy
    left_cell = header_table.cell(0, 0)
    left_cell.width = Cm(9.0)
    left_p = left_cell.paragraphs[0]
    run = left_p.add_run('ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ')
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = font_name
    for line in [
        'ΠΕΡΙΦΕΡΕΙΑ ΑΤΤΙΚΗΣ',
        'ΓΕΝΙΚΗ Δ/ΝΣΗ ΔΗΜΟΣΙΑΣ ΥΓΕΙΑΣ',
        '& ΚΟΙΝΩΝΙΚΗΣ ΜΕΡΙΜΝΑΣ',
        'Δ/ΝΣΗ ΚΟΙΝΩΝΙΚΗΣ ΜΕΡΙΜΝΑΣ',
        'ΤΜΗΜΑ ΚΟΙΝΩΝΙΚΗΣ ΑΛΛΗΛΕΓΓΥΗΣ',
    ]:
        run = left_p.add_run('\n' + line)
        run.font.size = Pt(8)
        run.font.name = font_name

    # Right column: date + protocol
    right_cell = header_table.cell(0, 1)
    right_cell.width = Cm(7.0)
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if not doc_date:
        from datetime import date as _date
        doc_date = _date.today().strftime('%d/%m/%Y')
    run = right_p.add_run(f'Αθήνα, {doc_date}')
    run.font.size = Pt(9)
    run.font.name = font_name
    if protocol_number:
        run = right_p.add_run(f'\nΑρ. Πρωτ.: {protocol_number}')
        run.font.size = Pt(9)
        run.font.name = font_name
    if internal_number:
        run = right_p.add_run(f'\nΕσωτ. Αρ.: {internal_number}')
        run.font.size = Pt(9)
        run.font.name = font_name

    doc.add_paragraph()  # spacer

    # ── 2. Title (ΘΕΜΑ) ───────────────────────────────────────
    theme_p = doc.add_paragraph()
    theme_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = theme_p.add_run(f'ΘΕΜΑ: {title}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = font_name

    doc.add_paragraph()  # spacer

    # ── 3. Legal references (Έχοντας υπόψη) ──────────────────
    if legal_references:
        ref_p = doc.add_paragraph()
        run = ref_p.add_run('Έχοντας υπόψη:')
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = font_name

        for i, ref in enumerate(legal_references, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(f'{i}. {ref}')
            run.font.size = Pt(9)
            run.font.name = font_name

        doc.add_paragraph()  # spacer

    # ── 4. Body text ──────────────────────────────────────────
    clean = re.sub(r'<br\s*/?>', '\n', rendered_html)
    clean = re.sub(r'<p>', '', clean)
    clean = re.sub(r'</p>', '\n\n', clean)
    clean = re.sub(r'<[^>]+>', '', clean)
    clean = html.unescape(clean)

    for para_text in clean.split('\n\n'):
        text = para_text.strip()
        if text:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = font_name

    # ── 5. Recipients table ───────────────────────────────────
    if recipients:
        doc.add_paragraph()  # spacer
        recip_p = doc.add_paragraph()
        recip_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = recip_p.add_run('ΠΙΝΑΚΑΣ ΑΠΟΔΕΚΤΩΝ')
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = font_name

        recip_table = doc.add_table(rows=len(recipients), cols=2)
        recip_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, r in enumerate(recipients):
            recip_table.cell(i, 0).text = str(i + 1) + '.'
            recip_table.cell(i, 1).text = r.get('name', '')
            for cell in recip_table.row_cells(i):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = font_name

    # Save to bytes
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
